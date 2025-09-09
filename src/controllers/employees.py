import base64
import io
import os
from email import encoders
from email.mime.application import MIMEApplication
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

import aiosmtplib
from decouple import config
from docx import Document
from openpyxl import load_workbook
from sqlalchemy import Integer, and_, cast, extract
from sqlmodel import select
from sqlmodel.ext.asyncio.session import AsyncSession

from src.config.database import engine
from src.helpers.decode_pdf_fields import decode_pdf_fields
from src.helpers.encode_pdf_fields import encode_pdf_fields
from src.helpers.ensure_pdf import ensure_pdf
from src.helpers.get_binary_fields import get_binary_fields
from src.helpers.handle_update_xlsx_rows import handle_update_xlsx_rows
from src.helpers.parse_dates import parse_dates
from src.helpers.set_checkbox import set_checkbox
from src.models.cities import Cities
from src.models.civil_status import CivilStatus
from src.models.cnh_categories import CnhCategories
from src.models.employees import Employees
from src.models.ethnicities import Ethnicities
from src.models.functions import Functions
from src.models.genders import Genders
from src.models.nationalities import Nationalities
from src.models.neighborhoods import Neighborhoods
from src.models.sectors import Sectors
from src.models.states import States
from src.models.subsidiaries import Subsidiaries
from src.models.turns import Turns
from src.schemas.employees import EmployeesBirthdayListProps, RowsListParams

SMTP_HOST = config("SMTP_HOST")

SMTP_PORT = config("SMTP_PORT")

SMTP_USER = config("SMTP_USER")

SMTP_PASS = config("SMTP_PASS")

SESI_EMAIL = config("SESI_EMAIL")

MABECON_EMAIL = config("MABECON_EMAIL")


async def handle_get_employees_by_subsidiarie(id: int):
    async with AsyncSession(engine) as session:
        result = await session.exec(
            select(Employees).where(Employees.subsidiarie_id == id)
        )

        employees = result.all()

        return [encode_pdf_fields(emp) for emp in employees]


async def handle_post_employees(employee: Employees):
    parse_dates(employee)

    decode_pdf_fields(employee)

    async with AsyncSession(engine) as session:
        session.add(employee)

        try:
            await session.commit()

            await session.refresh(employee)

            return employee

        except Exception as e:
            await session.rollback()

            raise e


async def handle_post_request_admissional_exam(id: int):
    async with AsyncSession(engine) as session:
        query = (
            select(Employees, Subsidiaries, Sectors, Functions)
            .join(Subsidiaries, Subsidiaries.id == Employees.subsidiarie_id)
            .join(Sectors, Sectors.id == Employees.sector_id)
            .join(Functions, Functions.id == Employees.function_id)
            .where(Employees.id == id)
        )

        db_result = await session.exec(query)

        result = db_result.first()

        if not result:
            raise ValueError("Funcionário não encontrado.")

        employee, subsidiarie, sector, function = result

        template_path = os.path.join("src", "assets", "exame_admissional_sesi.docx")

        doc = Document(template_path)

        placeholders = {
            "{{ subsidiarie }}": subsidiarie.name if subsidiarie else "",
            "{{ name }}": employee.name if employee else "",
            "{{ cpf }}": employee.cpf if employee else "",
            "{{ sector }}": sector.name if sector else "",
            "{{ function }}": function.name if function else "",
        }

        for p in doc.paragraphs:
            for ph, val in placeholders.items():
                if ph in p.text:
                    for run in p.runs:
                        if ph in run.text:
                            run.text = run.text.replace(ph, val)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for ph, val in placeholders.items():
                        if ph in cell.text:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    if ph in run.text:
                                        run.text = run.text.replace(ph, val)

        output_path = os.path.join(
            "src", "assets", f"admissional_exam_{employee.id}.docx"
        )

        doc.save(output_path)

        message = MIMEMultipart()

        message["From"] = SMTP_USER

        message["To"] = SESI_EMAIL

        message["Subject"] = f"Exame admissional de {employee.name}"

        body = f"""
        Segue em anexo solicitação de exame admissional para {employee.name}

        Atenciosamente,

        RH Postos Graciosa
        """

        message.attach(MIMEText(body, "plain"))

        with open(output_path, "rb") as f:
            part = MIMEBase("application", "octet-stream")

            part.set_payload(f.read())

            encoders.encode_base64(part)

            part.add_header(
                "Content-Disposition",
                f"attachment; filename={os.path.basename(output_path)}",
            )

            message.attach(part)

        await aiosmtplib.send(
            message,
            hostname=SMTP_HOST,
            port=SMTP_PORT,
            start_tls=True,
            username=SMTP_USER,
            password=SMTP_PASS,
        )

        os.remove(output_path)

        return {"status": "ok", "msg": f"E-mail enviado para {employee.email}"}


async def handle_post_send_employees_admission_to_contability(id: int):
    async with AsyncSession(engine) as session:
        template_path = "src/assets/ficha_da_contabilidade.xlsx"
        wb = load_workbook(template_path)
        ws = wb.active

        # Função helper para atualizar células apenas se o valor existir
        def safe_row(coord: str, value, attr: str = None):
            if value is None:
                return None
            if attr:
                v = getattr(value, attr, None)
                if v is None:
                    return None
                return RowsListParams(coord=coord, value=v)
            return RowsListParams(coord=coord, value=value)

        # Função helper para checkboxes seguros
        def set_safe_checkbox(start, end, value):
            if value is not None:
                set_checkbox(ws, start, end, value)

        # Busca funcionário
        employee = await session.get(Employees, id)
        if not employee:
            return {"error": f"Employee with id {id} not found"}

        # Busca todas as chaves estrangeiras manualmente
        employee_subsidiarie = (
            await session.get(Subsidiaries, employee.subsidiarie_id)
            if employee.subsidiarie_id
            else None
        )
        employee_gender = (
            await session.get(Genders, employee.gender_id)
            if employee.gender_id
            else None
        )
        employee_civil_status = (
            await session.get(CivilStatus, employee.civil_status_id)
            if employee.civil_status_id
            else None
        )
        employee_neighborhood = (
            await session.get(Neighborhoods, employee.neighborhood_id)
            if employee.neighborhood_id
            else None
        )
        employee_residence_city = (
            await session.get(Cities, employee.residence_city_id)
            if employee.residence_city_id
            else None
        )
        employee_ethnicitie = (
            await session.get(Ethnicities, employee.ethnicitie_id)
            if employee.ethnicitie_id
            else None
        )
        employee_birthcity = (
            await session.get(Cities, employee.birthcity_id)
            if employee.birthcity_id
            else None
        )
        employee_birthstate = (
            await session.get(States, employee.birthstate_id)
            if employee.birthstate_id
            else None
        )
        employee_nationalitie = (
            await session.get(Nationalities, employee.nationalitie_id)
            if employee.nationalitie_id
            else None
        )
        employee_rg_state = (
            await session.get(States, employee.rg_state_id)
            if employee.rg_state_id
            else None
        )
        employee_ctps_state = (
            await session.get(States, employee.ctps_state)
            if employee.ctps_state
            else None
        )
        employee_cnh_category = (
            await session.get(CnhCategories, employee.cnh_category_id)
            if employee.cnh_category_id
            else None
        )
        employee_function = (
            await session.get(Functions, employee.function_id)
            if employee.function_id
            else None
        )
        employee_residence_state = (
            await session.get(States, employee.residence_state_id)
            if employee.residence_state_id
            else None
        )
        employee_turn = (
            await session.get(Turns, employee.turn_id) if employee.turn_id else None
        )

        # Monta lista de rows com valores seguros
        rows_to_update = list(
            filter(
                None,
                [
                    safe_row("H1", employee_subsidiarie, "name"),
                    safe_row("H4", employee, "name"),
                    safe_row("AA4", employee_gender, "name"),
                    safe_row("AI4", employee_civil_status, "name"),
                    safe_row("J6", employee.street),
                    safe_row("Y6", employee.street_number),
                    safe_row("H7", employee_neighborhood, "name"),
                    safe_row("X7", employee.cep),
                    safe_row("AD7", employee_residence_city, "name"),
                    safe_row("H9", employee.phone),
                    safe_row("O9", employee.mobile),
                    safe_row("Y9", employee.email),
                    safe_row("AL9", employee_ethnicitie, "name"),
                    safe_row("R10", employee_birthcity, "name"),
                    safe_row("AB10", employee_birthstate, "name"),
                    safe_row("AJ10", employee_nationalitie, "name"),
                    safe_row("H11", employee.mothername),
                    safe_row("AF11", employee.fathername),
                    safe_row("H17", employee.cpf),
                    safe_row("H18", employee.rg),
                    safe_row("R18", employee.rg_issuing_agency),
                    safe_row("X18", employee_rg_state, "name"),
                    safe_row("R22", employee_ctps_state, "name"),
                    safe_row("AA18", employee.rg_expedition_date),
                    safe_row("H19", employee.military_certificate),
                    safe_row("H20", employee.pis),
                    safe_row("W20", employee.pis_register_date),
                    safe_row("H21", employee.votant_title),
                    safe_row("R21", employee.votant_zone),
                    safe_row("Y21", employee.votant_session),
                    safe_row("H22", employee.ctps),
                    safe_row("M22", employee.ctps_serie),
                    safe_row("T23", employee_cnh_category, "name"),
                    safe_row("H30", employee_function, "name"),
                    safe_row("Y22", employee.ctps_emission_date),
                    safe_row("H23", employee.cnh),
                    safe_row("Z23", employee.cnh_emission_date),
                    safe_row("AI23", employee.cnh_validity_date),
                    safe_row("W30", employee.admission_date),
                    safe_row("AD30", employee.monthly_wage),
                    safe_row("AJ30", employee.hourly_wage),
                    safe_row("AL30", employee.pro_rated_hours),
                    safe_row("AJ6", employee.street_complement),
                    safe_row("AM7", employee_residence_state, "name"),
                    safe_row("H10", employee.datebirth),
                    safe_row("AJ44", employee.ag),
                    safe_row("AM44", employee.cc),
                ],
            )
        )

        await handle_update_xlsx_rows(ws, rows_to_update)

        # Checkboxes seguros
        set_safe_checkbox("H25", "H26", employee.is_first_job)
        set_safe_checkbox("O25", "O26", employee.already_has_been_employee)
        set_safe_checkbox("W25", "W26", employee.trade_union_contribution_this_year)
        set_safe_checkbox("AA25", "AA26", employee.receiving_unemployment_insurance)
        set_safe_checkbox("AK25", "AK26", employee.has_previous_experience)
        set_safe_checkbox("P32", "M32", employee.has_harmfull_exposition)
        set_safe_checkbox("B35", "B36", employee.has_transport_voucher)
        set_safe_checkbox("K44", "H44", employee.has_hazard_pay)
        set_safe_checkbox("K45", "H45", employee.has_unhealthy_pay)

        # Subsidiária
        if employee_subsidiarie:
            ws["H46"] = "X" if employee_subsidiarie.id == 1 else employee_subsidiarie.id

        # Turnos
        if employee_turn:
            turn_map = {1: "J35", 2: "J36", 3: "J37"}
            cell = turn_map.get(employee.workdays_id)
            if cell:
                ws[cell] = "X"
                ws[f"R{cell[1:]}"] = employee_turn.start_time
                ws[f"X{cell[1:]}"] = employee_turn.start_interval_time
                ws[f"AA{cell[1:]}"] = employee_turn.end_interval_time
                ws[f"AF{cell[1:]}"] = employee_turn.end_time

        # Escolaridade
        school_map = {
            1: "AF18",
            2: "AF19",
            3: "AF20",
            4: "AK18",
            5: "AK19",
            6: "AK20",
            7: "AK21",
        }
        if employee.school_level_id:
            cell = school_map.get(employee.school_level_id)
            if cell:
                ws[cell] = "X"

        # Experiência
        if employee.experience_time_id:
            ws["B39"] = "X"
            ws["F39"] = "X"
            ws["H41"] = "X"
        else:
            ws["B40"] = "X"

        # Seguros e adiantamento salarial
        if employee.health_insurance:
            ws["W46"] = "X"
            ws["AB46"] = employee.health_insurance
        if employee.life_insurance:
            ws["AH46"] = "X"
            ws["AM46"] = employee.life_insurance
        if employee.wage_advance:
            ws["W45"] = "X"
            ws["AB45"] = employee.wage_advance

        # Pais
        columns = {
            "name": "J",
            "datebirth": "Y",
            "cityState": "AC",
            "cpf": "AG",
            "book": "AL",
            "paper": "AN",
        }
        for i, parent in enumerate(employee.parents[:3] if employee.parents else []):
            row = 13 + i
            for key, col in columns.items():
                ws[f"{col}{row}"] = parent.get(key, "")

        # Salva arquivo e cria email
        file_stream = io.BytesIO()
        wb.save(file_stream)
        file_stream.seek(0)
        message = MIMEMultipart()
        message["From"] = SMTP_USER
        message["To"] = MABECON_EMAIL
        message["Subject"] = (
            f"Encaminhamento de documentos do colaborador {employee.name} para admissão"
        )
        message.attach(
            MIMEText(
                f"Seguem em anexo os documentos do colaborador {employee.name} para admissão em {employee.admission_date}",
                "plain",
            )
        )

        file_name = f"ficha_da_contabilidade_{employee.name}.xlsx"
        part = MIMEApplication(file_stream.read(), Name=file_name)
        part["Content-Disposition"] = f'attachment; filename="{file_name}"'
        message.attach(part)

        # Anexos binários
        binary_fields = [
            ("ctps_file", f"ctps_{employee.name}.pdf"),
            (
                "admission_health_exam_file",
                f"exame_medico_admissional_{employee.name}.pdf",
            ),
            ("identity_file", f"identidade_{employee.name}.pdf"),
            ("cpf_file", f"cpf_{employee.name}.pdf"),
            ("votant_title_file", f"titulo_eleitor_{employee.name}.pdf"),
            ("residence_proof", f"comprovante_residencia_{employee.name}.pdf"),
            ("cnh_file", f"cnh_{employee.name}.pdf"),
            ("marriage_certificate_file", f"certificado_casamento_{employee.name}.pdf"),
            (
                "military_certificate_file",
                f"certificado_reservista_{employee.name}.pdf",
            ),
        ]
        for field, default_name in binary_fields:
            file_data = getattr(employee, field, None)
            if file_data:
                part = MIMEApplication(file_data, Name=default_name)
                part["Content-Disposition"] = f'attachment; filename="{default_name}"'
                message.attach(part)

        # Arquivos dos pais
        parent_files = [
            ("birthCertificate", "certidao_nascimento"),
            ("vaccinationCard", "carteira_vacinacao"),
            ("schoolingProof", "comprovante_escolar"),
        ]
        for i, parent in enumerate(employee.parents or []):
            for field, prefix in parent_files:
                file_b64 = parent.get(field)
                if file_b64:
                    try:
                        file_data = base64.b64decode(file_b64)
                        file_name = f"{prefix}_parent_{i + 1}.pdf"
                        part = MIMEApplication(file_data, Name=file_name)
                        part["Content-Disposition"] = (
                            f'attachment; filename="{file_name}"'
                        )
                        message.attach(part)
                    except Exception as e:
                        print(f"Erro ao decodificar {field} do parent {i + 1}: {e}")

        # Envio de email
        try:
            await aiosmtplib.send(
                message,
                hostname=SMTP_HOST,
                port=int(SMTP_PORT),
                username=SMTP_USER,
                password=SMTP_PASS,
                start_tls=True,
            )
        except Exception as e:
            return {"error": f"Falha ao enviar email: {str(e)}"}

        return {"status": "Email enviado com sucesso"}


async def handle_post_employees_birthday_list(body: EmployeesBirthdayListProps):
    async with AsyncSession(engine) as session:
        birthday_query = select(Employees).where(
            and_(
                Employees.subsidiarie_id == body.subsidiarie_id,
                extract("month", Employees.datebirth) == cast(body.month, Integer),
            )
        )

        query_result = await session.exec(birthday_query)

        employees_with_birthdays = query_result.all()

        return employees_with_birthdays


# NOTE: possível melhoria: só alterar o registro se for diferente do que está no banco real
# NOTE: testar primeiro, se falhar, implantar melhoria
async def handle_patch_employees(id: int, employee: Employees):
    parse_dates(employee)

    async with AsyncSession(engine) as session:
        query = select(Employees).where(Employees.id == id)

        result = await session.exec(query)

        db_employee = result.first()

        if not db_employee:
            return None

        for key, value in employee.dict(exclude_unset=True).items():
            if key in get_binary_fields():
                if isinstance(value, str):
                    try:
                        value = base64.b64decode(value)

                    except Exception as e:
                        raise ValueError(
                            f"Erro ao decodificar arquivo do campo {key}"
                        ) from e

                value = ensure_pdf(value, f"{key}.pdf")

            setattr(db_employee, key, value)

        session.add(db_employee)

        await session.commit()

        await session.refresh(db_employee)

        return encode_pdf_fields(db_employee)


# NOTE: depois employee.id vai ser chave estrangeira em outras tabelas
# NOTE: aí precisa mudar para delete on cascade
async def handle_delete_employees(id: int):
    async with AsyncSession(engine) as session:
        query = select(Employees).where(Employees.id == id)

        result = await session.exec(query)

        db_employee = result.first()

        if not db_employee:
            return None

        await session.delete(db_employee)

        await session.commit()

        return {"success": True}
